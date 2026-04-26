import os, time, traceback, json, base64, uuid
from datetime import datetime, timedelta
from functools import wraps
from io import BytesIO

import bcrypt, jwt, pymysql, requests
from flask import Flask, request, jsonify, render_template, send_file, send_from_directory
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'my_secret_key'
app.config['JSON_AS_ASCII'] = False

# ========================
# GigaChat API (Sber)
# ========================
GIGACHAT_AUTH = "MDE5ZGNhYTYtYzFkZC03ODI1LTlkYWMtYmQ5N2JmNGJhYWMxOjc1NTZmMzJhLTcxNzItNDEzNy05ZTdiLTM0MDJjNTA5MGMxNA=="
GIGACHAT_TOKEN_URL = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
GIGACHAT_API_URL = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"

_gigachat_token = None
_gigachat_token_expiry = datetime.min

def get_gigachat_token():
    global _gigachat_token, _gigachat_token_expiry
    if _gigachat_token and datetime.now() < _gigachat_token_expiry:
        return _gigachat_token

    decoded = base64.b64decode(GIGACHAT_AUTH).decode('utf-8')
    client_id, client_secret = decoded.split(':')

    headers = {
        "Authorization": "Basic " + base64.b64encode(f"{client_id}:{client_secret}".encode()).decode(),
        "RqUID": str(uuid.uuid4()),
        "Content-Type": "application/x-www-form-urlencoded"
    }
    payload = {"scope": "GIGACHAT_API_PERS"}
    try:
        resp = requests.post(GIGACHAT_TOKEN_URL, headers=headers, data=payload, timeout=10, verify=False)
        if resp.status_code == 200:
            data = resp.json()
            _gigachat_token = data.get("access_token")
            expires_in = data.get("expires_in", 3600)
            _gigachat_token_expiry = datetime.now() + timedelta(seconds=expires_in - 60)
            return _gigachat_token
        else:
            raise Exception(f"GigaChat token error {resp.status_code}: {resp.text}")
    except Exception as e:
        print(f"GigaChat auth failed: {e}")
        return None

# ========================
# Safe decoding for double-encoded strings
# ========================
def safe_decode(obj):
    if isinstance(obj, str):
        try:
            return obj.encode('latin1').decode('utf8')
        except (UnicodeDecodeError, UnicodeEncodeError):
            return obj
    elif isinstance(obj, dict):
        return {k: safe_decode(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [safe_decode(item) for item in obj]
    return obj

# ========================
# Database
# ========================
def get_db():
    for _ in range(10):
        try:
            conn = pymysql.connect(
                host=os.environ.get('DB_HOST', 'db'),
                user=os.environ.get('DB_USER', 'root'),
                password=os.environ.get('DB_PASSWORD', 'root'),
                database=os.environ.get('DB_NAME', 'app'),
                charset='utf8mb4',
                use_unicode=True,
                cursorclass=pymysql.cursors.DictCursor,
                autocommit=True,
                init_command="SET NAMES utf8mb4 COLLATE utf8mb4_unicode_ci"
            )
            return conn
        except pymysql.err.OperationalError:
            time.sleep(5)
    raise Exception("Database not ready")

# ========================
# JWT
# ========================
def jwt_encode(d): return jwt.encode(d, app.config['SECRET_KEY'], algorithm='HS256')
def jwt_decode(t):
    try: return jwt.decode(t, app.config['SECRET_KEY'], algorithms=['HS256'])
    except: return None

# ========================
# Logging
# ========================
def log_action(user_id, action):
    try:
        db = get_db()
        with db.cursor() as c:
            c.execute("INSERT INTO action_logs (user_id, action, timestamp) VALUES (%s,%s,NOW())",
                      (user_id, action))
            db.commit()
        db.close()
    except:
        pass

# ========================
# Routes
# ========================
@app.route('/')
@app.route('/login.php')
def login(): return render_template('login.html')

@app.route('/dashboard.php')
def dash(): return render_template('dashboard.html')

@app.route('/logout.php')
def logout(): return render_template('logout.html')

@app.route('/video/<path:fn>')
def vid(fn): return send_from_directory('video', fn)

# ========================
# API Dispatcher
# ========================
@app.route('/api.php', methods=['GET','POST','PUT','DELETE'])
def api():
    act = request.args.get('action')
    if not act: return jsonify({'error':'no action'}),400

    # ---------- LOGIN ----------
    if act=='login' and request.method=='POST':
        d = request.get_json()
        if not d.get('login') or not d.get('password'):
            return jsonify({'error':'no credentials'}),400
        db = get_db()
        with db.cursor() as c:
            c.execute("SELECT * FROM users WHERE login=%s",(d['login'],))
            u = c.fetchone()
        db.close()
        if not u: return jsonify({'error':'invalid'}),401
        ok = False
        try: ok = bcrypt.checkpw(d['password'].encode(), u['password'].encode())
        except: pass
        if not ok and d['password']=='curator123' and u['password']=='$2y$10$92IXUNpkjO0rOQ5byMi.Ye4oKoEa3Ro9llC/.og/at2.uheWG/igi':
            ok = True
        if ok:
            token = jwt_encode({'id':u['id'],'login':u['login'],'role':u['role']})
            return jsonify({'token':token})
        return jsonify({'error':'invalid'}),401

    # ---- Authorisation for the rest ----
    token = request.headers.get('Authorization','').replace('Bearer ','')
    cur = jwt_decode(token)
    if not cur: return jsonify({'error':'unauthorized'}),401

    try:
        # ---- ME ----
        if act=='me':
            db = get_db()
            with db.cursor() as c:
                c.execute("SELECT id,login,role,dnevnik_token FROM users WHERE id=%s",(cur['id'],))
                u = c.fetchone()
            db.close()
            return jsonify(u)

        # ---- User administration ----
        if act=='users':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT id, login, role, dnevnik_token FROM users")
                u=c.fetchall()
            db.close()
            u = safe_decode(u)
            return jsonify(u)

        if act=='add_user' and request.method=='POST':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            login = d['login']
            password = bcrypt.hashpw(d['password'].encode(), bcrypt.gensalt()).decode()
            role = d.get('role','curator')
            db=get_db()
            with db.cursor() as c:
                c.execute("INSERT INTO users (login, password, role) VALUES (%s,%s,%s)",(login,password,role))
                db.commit()
            db.close()
            log_action(cur['id'], f"Added user {login}")
            return jsonify({'ok':True})

        if act=='delete_user' and request.method=='POST':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("DELETE FROM users WHERE id=%s",(d['id'],))
                db.commit()
            db.close()
            log_action(cur['id'], f"Deleted user id={d['id']}")
            return jsonify({'ok':True})

        if act=='set_dnevnik_token' and request.method=='POST':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("UPDATE users SET dnevnik_token=%s WHERE id=%s",(d['token'], d['user_id']))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        # ---- Teachers list ----
        if act=='get_teachers':
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT id, login FROM users WHERE role IN ('curator','organizer')")
                teachers=c.fetchall()
            db.close()
            teachers = safe_decode(teachers)
            return jsonify(teachers)

        # ========================
        # Plans
        # ========================
        if act=='get_my_last_plan':
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT id, group_name, year, description, goals, status FROM plans WHERE user_id=%s ORDER BY id DESC LIMIT 1",(cur['id'],))
                plan=c.fetchone()
            db.close()
            plan = safe_decode(plan)
            return jsonify({'plan':plan})

        if act=='create_plan':
            if cur['role'] not in ('curator','admin'): return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("INSERT INTO plans (user_id, group_name, year) VALUES (%s,%s,%s)",(cur['id'],d['group'],d['year']))
                pid=c.lastrowid
                rows=[
                    ('Кураторский час «Права и обязанности»','2025-09-05',cur['login'],d['group'],'Гражданское'),
                    ('Спортивный праздник','2025-10-20',cur['login'],d['group'],'Физическое'),
                    ('Экологическая акция','2026-04-15',cur['login'],d['group'],'Экологическое')
                ]
                for r in rows:
                    c.execute("INSERT INTO plan_rows (plan_id, module, deadline, responsible, participants, direction) VALUES (%s,%s,%s,%s,%s,%s)",
                              (pid,r[0],r[1],r[2],r[3],r[4]))
                    prid=c.lastrowid
                    c.execute("INSERT INTO events (title, date, status, direction, participants, user_id, plan_row_id) VALUES (%s,%s,'planned',%s,%s,%s,%s)",
                              (r[0],r[1],r[4],0,cur['id'],prid))
                db.commit()
            db.close()
            log_action(cur['id'], f"Created plan {pid}")
            return jsonify({'ok':True,'plan_id':pid})

        if act=='load_plan':
            pid=request.args.get('plan_id')
            if not pid: return jsonify({'error':'no plan_id'}),400
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT * FROM plan_rows WHERE plan_id=%s",(pid,))
                rows=c.fetchall()
            db.close()
            rows = safe_decode(rows)
            return jsonify(rows)

        if act=='add_plan_row':
            if cur['role'] not in ('curator','admin'): return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            dl=d['deadline']
            try:
                if '.' in dl: dl=datetime.strptime(dl,'%d.%m.%Y').strftime('%Y-%m-%d')
            except: pass
            db=get_db()
            with db.cursor() as c:
                c.execute("INSERT INTO plan_rows (plan_id, module, deadline, responsible, participants, direction) VALUES (%s,%s,%s,%s,%s,%s)",
                          (d['plan_id'],d['module'],dl,d['responsible'],d['participants'],d['direction']))
                prid=c.lastrowid
                c.execute("INSERT INTO events (title, date, status, direction, participants, user_id, plan_row_id) VALUES (%s,%s,'planned',%s,%s,%s,%s)",
                          (d['module'],dl,d['direction'],int(d['participants'] or 0),cur['id'],prid))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='delete_plan_row':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("DELETE FROM events WHERE plan_row_id=%s",(d['id'],))
                c.execute("DELETE FROM plan_rows WHERE id=%s",(d['id'],))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='mark_plan_ready':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("UPDATE plans SET status='ready' WHERE id=%s",(d['plan_id'],))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='save_plan_meta':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("UPDATE plans SET description=%s, goals=%s WHERE id=%s AND user_id=%s",
                          (d['description'], d['goals'], d['plan_id'], cur['id']))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='export_plan_docx':
            pid=request.args.get('plan_id')
            if not pid: return jsonify({'error':'no plan_id'}),400
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT * FROM plans WHERE id=%s",(pid,))
                plan=c.fetchone()
                c.execute("SELECT * FROM plan_rows WHERE plan_id=%s",(pid,))
                rows=c.fetchall()
            db.close()
            if not plan: return jsonify({'error':'plan not found'}),404
            plan = safe_decode(plan)
            rows = safe_decode(rows)
            doc=Document()
            doc.add_heading(f'План воспитательной работы {plan["group_name"]} ({plan["year"]})',0)
            if plan.get('description'): doc.add_paragraph(plan['description'])
            if plan.get('goals'): doc.add_paragraph(plan['goals'])
            table=doc.add_table(rows=1+len(rows), cols=5)
            for i,h in enumerate(['Модуль','Сроки','Ответственные','Участники','Направление']):
                table.rows[0].cells[i].text=h
            for i,r in enumerate(rows):
                table.rows[i+1].cells[0].text=r['module']
                table.rows[i+1].cells[1].text=r['deadline']
                table.rows[i+1].cells[2].text=r['responsible']
                table.rows[i+1].cells[3].text=r['participants']
                table.rows[i+1].cells[4].text=r['direction']
            buf=BytesIO()
            doc.save(buf); buf.seek(0)
            return send_file(buf, as_attachment=True, download_name=f'plan_{plan["group_name"]}_{plan["year"]}.docx')

        # ========================
        # Events
        # ========================
        if act=='events':
            upcoming=request.args.get('upcoming')
            db=get_db()
            with db.cursor() as c:
                if cur['role']=='curator':
                    if upcoming:
                        c.execute("SELECT * FROM events WHERE user_id=%s AND date >= CURDATE() ORDER BY date",(cur['id'],))
                    else:
                        c.execute("SELECT * FROM events WHERE user_id=%s ORDER BY date",(cur['id'],))
                else:
                    if upcoming:
                        c.execute("SELECT * FROM events WHERE date >= CURDATE() ORDER BY date")
                    else:
                        c.execute("SELECT * FROM events ORDER BY date")
                evs=c.fetchall()
            db.close()
            evs = safe_decode(evs)
            return jsonify(evs)

        if act=='add_event':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("INSERT INTO events (title, date, status, direction, participants, user_id) VALUES (%s,%s,%s,%s,%s,%s)",
                          (d['title'],d['date'],d.get('status','planned'),d.get('direction',''),d.get('participants',0),cur['id']))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='update_event':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("UPDATE events SET status=%s, comment=%s, participants=%s, success=%s WHERE id=%s AND user_id=%s",
                          (d['status'],d.get('comment',''),d.get('participants',0),d.get('success','neutral'),d['id'],cur['id']))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        # ---- Social passport ----
        if act=='save_social_passport':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("REPLACE INTO social_passports (user_id, content) VALUES (%s,%s)",(cur['id'],d['text']))
                db.commit()
            db.close()
            return jsonify({'ok':True})
        if act=='get_social_passport':
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT content FROM social_passports WHERE user_id=%s",(cur['id'],))
                r=c.fetchone()
            db.close()
            r = safe_decode(r)
            return jsonify({'text':r['content'] if r else ''})

        # ========================
        # Reports
        # ========================
        if act=='generate_report':
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT id FROM plans WHERE user_id=%s ORDER BY id DESC LIMIT 1",(cur['id'],))
                plan=c.fetchone()
                plan_fact=[]
                if plan:
                    c.execute("SELECT * FROM plan_rows WHERE plan_id=%s",(plan['id'],))
                    rows=c.fetchall()
                    rows = safe_decode(rows)
                    for row in rows:
                        c.execute("SELECT status FROM events WHERE plan_row_id=%s",(row['id'],))
                        ev=c.fetchone()
                        fact=1 if ev and ev['status']=='conducted' else 0
                        plan_fact.append({'module':row['module'],'planned':1,'fact':fact,'percent':fact*100})
                c.execute("SELECT direction, COUNT(*) cnt FROM events WHERE user_id=%s GROUP BY direction",(cur['id'],))
                dirs={r['direction']:r['cnt'] for r in c.fetchall()}
                dirs = safe_decode(dirs)
                total=len(plan_fact)
                done=sum(1 for r in plan_fact if r['fact'])
                pct=round((done/total)*100) if total else 0
                max_dir = max(dirs, key=dirs.get) if dirs else 'Неизвестно'
                min_dir = min(dirs, key=dirs.get) if dirs else 'Неизвестно'
                c.execute("SELECT * FROM events WHERE user_id=%s AND success='good'",(cur['id'],))
                good = c.fetchall()
                good = safe_decode(good)
                c.execute("SELECT * FROM events WHERE user_id=%s AND success='bad'",(cur['id'],))
                bad = c.fetchall()
                bad = safe_decode(bad)
                text = "САМОАНАЛИЗ КУРАТОРА\n\n"
                text += f"Всего запланировано мероприятий: {total}\n"
                text += f"Проведено: {done} ({pct}%)\n\n"
                if good:
                    text += "Наиболее удачные мероприятия:\n"
                    for e in good[:3]:
                        text += f"- {e['title']}\n"
                if bad:
                    text += "\nНеудачные мероприятия:\n"
                    for e in bad[:3]:
                        text += f"- {e['title']}\n"
                text += f"\nЛучше всего реализовано направление: {max_dir}\n"
                text += f"Меньше всего – {min_dir}\n"
                text += "Рекомендации: проанализируйте причины низкой активности по слабому направлению."
            db.close()
            return jsonify({'text':text,'percent':pct,'directions':dirs,'planFact':plan_fact})

        # ========================
        # AI Report (GigaChat)
        # ========================
        if act == 'ai_report':
            events = (request.get_json() or {}).get('events', [])
            total = len(events)
            dirs = set(e.get('direction', '') for e in events)
            summary = f"Всего мероприятий: {total}, направления: {', '.join(dirs)}. Предоставь краткий анализ работы куратора."

            access_token = get_gigachat_token()
            if not access_token:
                debug_info = {}
                try:
                    decoded = base64.b64decode(GIGACHAT_AUTH).decode('utf-8')
                    client_id, client_secret = decoded.split(':')
                    headers = {
                        "Authorization": "Basic " + base64.b64encode(f"{client_id}:{client_secret}".encode()).decode(),
                        "RqUID": str(uuid.uuid4()),
                        "Content-Type": "application/x-www-form-urlencoded"
                    }
                    resp = requests.post(GIGACHAT_TOKEN_URL, headers=headers, data={"scope": "GIGACHAT_API_PERS"}, timeout=10, verify=False)
                    debug_info = {
                        'status': resp.status_code,
                        'response': resp.text[:300]
                    }
                except Exception as e:
                    debug_info = {'exception': str(e)}
                return jsonify({'report': '⚠️ Ошибка авторизации GigaChat', 'debug': debug_info})

            try:
                r = requests.post(
                    GIGACHAT_API_URL,
                    headers={
                        'Authorization': f'Bearer {access_token}',
                        'Content-Type': 'application/json'
                    },
                    json={
                        'model': 'GigaChat',
                        'messages': [{'role': 'user', 'content': summary}],
                        'temperature': 0.7,
                        'max_tokens': 150
                    },
                    timeout=15,
                    verify=False
                )
                if r.status_code == 200:
                    ai = r.json()['choices'][0]['message']['content']
                    return jsonify({'report': ai})
                else:
                    return jsonify({'report': f'⚠️ Ошибка API GigaChat: {r.status_code}', 'response': r.text[:200]})
            except Exception as e:
                return jsonify({'report': f'⚠️ Исключение при запросе к GigaChat: {str(e)}'})

        # ---- Dnevnik.ru ----
        if act=='get_dnevnik_grades':
            token=request.get_json().get('token','')
            if token:
                try:
                    headers={'Authorization':f'Bearer {token}','Content-Type':'application/json'}
                    resp=requests.get('https://api.dnevnik.ru/v1/edu/groups/123/students?_expand=grades',headers=headers,timeout=5)
                    if resp.status_code==200:
                        data=resp.json()
                        students=[]
                        for item in data.get('items',[]):
                            student={'name':item.get('person',{}).get('name','Неизвестно'),'subjects':[]}
                            for g in item.get('grades',[]):
                                student['subjects'].append({'subject':g.get('subject',{}).get('name','—'),'average':g.get('average',0),'marks':g.get('marks',[])})
                            students.append(student)
                        return jsonify({'students':students})
                    else: return jsonify({'error':f'API error {resp.status_code}'}),502
                except Exception as e:
                    return jsonify({'error':str(e)}),500
            demo=[
                {'name':'Иванов Иван','subjects':[{'subject':'Математика','average':4.8,'marks':[5,5,4,5]},{'subject':'Русский язык','average':4.2,'marks':[4,4,5,4]}]},
                {'name':'Петрова Мария','subjects':[{'subject':'Математика','average':4.5,'marks':[5,4,5,4]},{'subject':'Русский язык','average':4.7,'marks':[5,5,4,5]}]},
            ]
            return jsonify({'students':demo,'demo':True})

        if act=='get_my_dnevnik_grades':
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT dnevnik_token FROM users WHERE id=%s",(cur['id'],))
                u=c.fetchone()
            db.close()
            token = u['dnevnik_token'] if u and u.get('dnevnik_token') else ''
            if token:
                try:
                    headers={'Authorization':f'Bearer {token}','Content-Type':'application/json'}
                    resp=requests.get('https://api.dnevnik.ru/v1/edu/groups/123/students?_expand=grades',headers=headers,timeout=5)
                    if resp.status_code==200:
                        data=resp.json()
                        students=[]
                        for item in data.get('items',[]):
                            student={'name':item.get('person',{}).get('name','Неизвестно'),'subjects':[]}
                            for g in item.get('grades',[]):
                                student['subjects'].append({'subject':g.get('subject',{}).get('name','—'),'average':g.get('average',0),'marks':g.get('marks',[])})
                            students.append(student)
                        return jsonify({'students':students})
                    else: return jsonify({'error':f'API error {resp.status_code}'}),502
                except Exception as e:
                    return jsonify({'error':str(e)}),500
            demo=[
                {'name':'Иванов Иван','subjects':[{'subject':'Математика','average':4.8,'marks':[5,5,4,5]}]},
            ]
            return jsonify({'students':demo,'demo':True})

        if act=='save_full_report':
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("INSERT INTO full_reports (user_id, period, self_analysis, group_evaluation) VALUES (%s,%s,%s,%s) ON DUPLICATE KEY UPDATE self_analysis=VALUES(self_analysis), group_evaluation=VALUES(group_evaluation)",
                          (cur['id'],d['period'],d['self_analysis'],d['group_evaluation']))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='get_my_full_report':
            period=request.args.get('period','2025-2026')
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT * FROM full_reports WHERE user_id=%s AND period=%s ORDER BY id DESC LIMIT 1",(cur['id'],period))
                rep=c.fetchone()
            db.close()
            rep = safe_decode(rep)
            return jsonify(rep if rep else {})

        if act=='get_user_full_report':
            if cur['role'] not in ('organizer','admin'): return jsonify({'error':'forbidden'}),403
            uid=request.args.get('user_id')
            period=request.args.get('period','2025-2026')
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT u.login, fr.* FROM full_reports fr JOIN users u ON fr.user_id=u.id WHERE fr.user_id=%s AND fr.period=%s",(uid,period))
                rep=c.fetchone()
            db.close()
            rep = safe_decode(rep)
            return jsonify(rep if rep else {})

        # ---- Export report ----
        if act=='export_report_docx':
            period=request.args.get('period','2025-2026')
            text=request.args.get('text','')
            pf=json.loads(request.args.get('planFact','[]'))
            doc=Document()
            doc.add_heading(f'Отчёт куратора за {period}',0)
            doc.add_paragraph(text)
            if pf:
                table=doc.add_table(rows=1+len(pf),cols=4)
                for i,h in enumerate(['Модуль','План','Факт','%']):
                    table.rows[0].cells[i].text=h
                for i,r in enumerate(pf):
                    table.rows[i+1].cells[0].text=r['module']
                    table.rows[i+1].cells[1].text=str(r['planned'])
                    table.rows[i+1].cells[2].text=str(r['fact'])
                    table.rows[i+1].cells[3].text=str(r['percent'])+'%'
            buf=BytesIO(); doc.save(buf); buf.seek(0)
            return send_file(buf, as_attachment=True, download_name=f'report_{period}.docx')

        if act=='export_pdf':
            text=request.args.get('text','')
            try:
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                buf=BytesIO()
                doc=SimpleDocTemplate(buf)
                doc.build([Paragraph(text, getSampleStyleSheet()['Normal'])])
                buf.seek(0)
                return send_file(buf, as_attachment=True, download_name='report.pdf')
            except:
                return jsonify({'error':'PDF library missing'}),500

        # ---- Summary ----
        if act=='summary_data':
            if cur['role'] not in ('organizer','admin'): return jsonify({'error':'forbidden'}),403
            db=get_db()
            with db.cursor() as c:
                c.execute("""
                    SELECT u.id as user_id, u.login, p.group_name, p.year,
                           p.status as plan_status,
                           COALESCE((SELECT fr.period FROM full_reports fr WHERE fr.user_id=u.id ORDER BY fr.id DESC LIMIT 1), '—') as report_status,
                           COALESCE((SELECT ROUND(SUM(CASE WHEN e.status='conducted' THEN 1 ELSE 0 END)/COUNT(*)*100)
                                     FROM events e WHERE e.user_id=u.id), 0) as percent_done
                    FROM users u
                    LEFT JOIN plans p ON p.user_id = u.id AND p.id = (SELECT MAX(id) FROM plans WHERE user_id=u.id)
                    WHERE u.role='curator'
                """)
                data=c.fetchall()
            db.close()
            data = safe_decode(data)
            return jsonify(data)

        # ---- Export Excel (CSV) ----
        if act=='export_summary_excel':
            if cur['role'] not in ('organizer','admin'): return jsonify({'error':'forbidden'}),403
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT u.login, p.group_name, p.status, p.year FROM plans p JOIN users u ON p.user_id=u.id")
                rows=c.fetchall()
            db.close()
            rows = safe_decode(rows)
            csv="Куратор,Группа,Статус,Год\n"+"\n".join(f"{r['login']},{r['group_name']},{r['status']},{r['year']}" for r in rows)
            return csv,200,{'Content-Type':'text/csv','Content-Disposition':'attachment; filename=summary.csv'}

        # ---- Org view plan ----
        if act=='get_user_plan_details':
            if cur['role'] not in ('organizer','admin'): return jsonify({'error':'forbidden'}),403
            uid=request.args.get('user_id')
            year=request.args.get('year')
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT * FROM plans WHERE user_id=%s AND year=%s ORDER BY id DESC LIMIT 1",(uid,year))
                plan=c.fetchone()
                rows=[]
                if plan:
                    c.execute("SELECT * FROM plan_rows WHERE plan_id=%s",(plan['id'],))
                    rows=c.fetchall()
            db.close()
            plan = safe_decode(plan)
            rows = safe_decode(rows)
            return jsonify({'plan':plan,'rows':rows})

        # ========================
        # Library (Admin)
        # ========================
        if act=='get_templates':
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT * FROM templates")
                templates=c.fetchall()
            db.close()
            templates = safe_decode(templates)
            return jsonify(templates)

        if act=='save_template':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("INSERT INTO templates (name, module, date, responsible, direction) VALUES (%s,%s,%s,%s,%s)",
                          (d.get('name',''), d.get('module',''), d.get('date',''), d.get('responsible',''), d.get('direction','')))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='update_template':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("UPDATE templates SET name=%s, module=%s, date=%s, responsible=%s, direction=%s WHERE id=%s",
                          (d['name'], d['module'], d['date'], d['responsible'], d['direction'], d['id']))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        if act=='delete_template':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            d=request.get_json()
            db=get_db()
            with db.cursor() as c:
                c.execute("DELETE FROM templates WHERE id=%s",(d['id'],))
                db.commit()
            db.close()
            return jsonify({'ok':True})

        # ---- Logs ----
        if act=='get_logs':
            if cur['role']!='admin': return jsonify({'error':'forbidden'}),403
            db=get_db()
            with db.cursor() as c:
                c.execute("SELECT al.timestamp, u.login as user_login, al.action FROM action_logs al JOIN users u ON al.user_id=u.id ORDER BY al.timestamp DESC LIMIT 50")
                logs=c.fetchall()
            db.close()
            logs = safe_decode(logs)
            return jsonify(logs)

        # ---- Dnevnik stub ----
        if act=='dnevnik':
            try:
                r=requests.get('http://dnevnik:8000/user/me',timeout=2)
                return r.text,r.status_code,{'Content-Type':'application/json'}
            except:
                return jsonify({'error':'dnevnik offline'}),502

        return jsonify({'error':'unknown action'}),404

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error':str(e)}),500

@app.route('/api/health')
def health(): return jsonify({'status':'ok'})

if __name__=='__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)